from dataclasses import dataclass
from pathlib import Path
import re


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".html", ".htm"}


@dataclass(frozen=True)
class ParsedDocument:
    parser_name: str
    text: str
    metadata: dict


def parse_file(path: Path, file_ext: str) -> ParsedDocument:
    ext = file_ext.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".docx":
        return parse_docx(path)
    if ext in {".md", ".markdown"}:
        return parse_markdown(path)
    if ext in {".html", ".htm"}:
        return parse_html(path)
    raise ValueError(f"Unsupported document type: {file_ext}")


def parse_pdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    metadata = reader.metadata or {}
    pages: list[str] = []
    empty_pages: list[int] = []

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise ValueError("PDF is encrypted and cannot be opened without a password") from exc

    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if not page_text.strip():
            empty_pages.append(index)
        pages.append(f"[Page {index}]\n{page_text.strip()}")

    text = _normalize_text("\n\n".join(pages))
    return ParsedDocument(
        parser_name="pypdf",
        text=text,
        metadata={
            "page_count": len(reader.pages),
            "empty_pages": empty_pages,
            "title": _safe_metadata_value(metadata.get("/Title")),
            "author": _safe_metadata_value(metadata.get("/Author")),
            "producer": _safe_metadata_value(metadata.get("/Producer")),
        },
    )


def parse_docx(path: Path) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    blocks: list[str] = []
    headings: list[dict] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            headings.append({"style": style_name, "text": text})
            blocks.append(f"{text}")
        else:
            blocks.append(text)

    table_count = 0
    for table in doc.tables:
        table_count += 1
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            blocks.append("\n".join(rows))

    props = doc.core_properties
    text = _normalize_text("\n\n".join(blocks))
    return ParsedDocument(
        parser_name="python-docx",
        text=text,
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": table_count,
            "headings": headings,
            "title": _safe_metadata_value(props.title),
            "author": _safe_metadata_value(props.author),
            "created": _safe_metadata_value(props.created),
            "modified": _safe_metadata_value(props.modified),
        },
    )


def parse_markdown(path: Path) -> ParsedDocument:
    raw = _read_text(path)
    headings = []
    for line_no, line in enumerate(raw.splitlines(), start=1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            headings.append(
                {
                    "line": line_no,
                    "level": len(match.group(1)),
                    "text": match.group(2),
                }
            )

    frontmatter = raw.startswith("---\n") or raw.startswith("---\r\n")
    return ParsedDocument(
        parser_name="markdown-text",
        text=_normalize_text(raw),
        metadata={
            "heading_count": len(headings),
            "headings": headings,
            "frontmatter": frontmatter,
            "link_count": len(re.findall(r"\[[^\]]+\]\([^)]+\)", raw)),
        },
    )


def parse_html(path: Path) -> ParsedDocument:
    from bs4 import BeautifulSoup

    raw = _read_text(path)
    soup = BeautifulSoup(raw, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    title = soup.title.get_text(strip=True) if soup.title else None
    description_tag = soup.find("meta", attrs={"name": "description"})
    description = description_tag.get("content") if description_tag else None
    headings = [
        {"level": int(tag.name[1]), "text": tag.get_text(" ", strip=True)}
        for tag in soup.find_all(re.compile("^h[1-6]$"))
    ]
    text = _normalize_text(soup.get_text("\n", strip=True))

    return ParsedDocument(
        parser_name="beautifulsoup4",
        text=text,
        metadata={
            "title": title,
            "description": description,
            "heading_count": len(headings),
            "headings": headings,
        },
    )


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    compact = "\n".join(lines)
    return re.sub(r"\n{4,}", "\n\n\n", compact).strip()


def _safe_metadata_value(value: object) -> str | None:
    if value is None:
        return None
    return str(value)

